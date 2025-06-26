
from flask import Flask, request, send_file
from docx import Document
from io import BytesIO
from flask_cors import CORS

app = Flask(__name__)
CORS(app)

@app.route('/generate', methods=['POST'])
def generate_opex():
    data = request.get_json()
    print("Received OPEX data:", data)

    try:
        selected_years = data.get("template", "").split(",")
        selected_years = [int(y.strip()) for y in selected_years if y.strip().isdigit()]
        if not selected_years:
            return {"error": "No valid template selected"}, 400

        # Compute generation
        gen_number = float(data.get("generationNumber", "0").replace(',', ''))
        capacity = float(data.get("capacity", "0").replace(',', ''))
        gen1 = gen_number * capacity
        data["generationFirstYear"] = round(gen1)

        expected, guaranteed = [], []
        current = gen1
        for i in range(25):
            if i == 1:
                current *= (1 - 0.02)
            elif i > 1:
                current *= (1 - 0.0055)
            expected.append(round(current))
            guaranteed.append(round(current * 0.9))
            data[f"expectedGen{i+1}"] = expected[i]
            data[f"guaranteedGen{i+1}"] = guaranteed[i]

        # Build savings logic if applicable
        include_savings = data.get("include_savings", "no").lower() == "yes"
        template_key = data["template"].replace(",", "^J") + ("-savings" if include_savings else "")
        template_filename = f"proposal-temp-{template_key}.docx"
        doc = Document(f"savings_templates/{template_filename}")

        if include_savings:
            current_price = float(data.get("current_electricity_price", "0").replace(',', ''))
            sum_savings = 0
            for i in range(1, 26):
                grid_price = current_price * ((1.015) ** (i - 1))
                expected_gen = expected[i - 1]
                year = str(i)
                grid_bill = grid_price * expected_gen
                data[f"current_electricity_price-{year}"] = round(grid_price, 2)
                data[f"Annual_Grid_Bill-{year}"] = round(grid_bill)
                # Tariff for the first selected year
                first_year = selected_years[0]
                tariff = float(data.get(f"tariff_{first_year}", "0"))
                solar_bill = tariff * expected_gen
                data[f"Annual_Solar_Bill-{year}"] = round(solar_bill)
                data[f"Annual_Cost_Savings-{year}"] = round(grid_bill - solar_bill)
                if i <= max(selected_years):
                    sum_savings += grid_bill - solar_bill
            data[f"Sum_Cost_Savings-{max(selected_years)}"] = round(sum_savings)

        # Termination and tariff for each selected year
        for year in selected_years:
            data[f"tariff-{year}"] = data.get(f"tariff_{year}", "0")
            terms = data.get(f"termination_{year}", "")
            term_lines = [x.strip().replace(',', '') for x in terms.split("\n") if x.strip()]
            for i in range(year):
                data[f"termPayment-{year}-{i+1}"] = term_lines[i] if i < len(term_lines) else "0"

        # Text replacement logic
        def replace_all_placeholders(doc_obj, replacements):
            def replace_text(text):
                for key, val in replacements.items():
                    text = text.replace(f"{{{{{key}}}}}", str(val))
                return text
            for p in doc_obj.paragraphs:
                full_text = ''.join(run.text for run in p.runs)
                if any(f"{{{{{k}}}}}" in full_text for k in replacements):
                    new_text = replace_text(full_text)
                    for run in p.runs:
                        run.text = ""
                    p.runs[0].text = new_text
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_all_placeholders(cell, replacements)

        replace_all_placeholders(doc, data)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"Cleanmax Proposal-{data.get('clientName', 'Client')}.docx"
        return send_file(output, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        import traceback
        traceback.print_exc()
        return {"error": str(e)}, 500

if __name__ == '__main__':
    app.run(debug=True)
