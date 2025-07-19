
from flask import Flask, request, send_file, jsonify
from docx import Document
from io import BytesIO
from flask_cors import CORS
import re
import os
import tempfile
import pythoncom
import win32com.client
import threading

app = Flask(__name__)
CORS(app)

def calculate_npv(values, rate):
    return sum(v / ((1 + rate) ** (i + 1)) for i, v in enumerate(values))

def format_inr(n):
    try:
        n = int(round(n))
        s = str(n)
        if len(s) <= 3:
            return s
        last3 = s[-3:]
        rest = s[:-3]
        rest = re.sub(r'(\d)(?=(\d{2})+(?!\d))', r'\1,', rest)
        return rest + ',' + last3
    except:
        return str(n)

def replace_all_placeholders(doc_obj, replacements):
    def replace_text(text):
        for key, val in replacements.items():
            text = text.replace(f"{{{{{key}}}}}", str(val))
        return text

    for p in doc_obj.paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        if any(f"{{{{{k}}}}}" in full_text for k in replacements):
            new_text = replace_text(full_text)
            for run in p.runs:
                run.text = ""
            p.runs[0].text = new_text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_all_placeholders(cell, replacements)

def convert_docx_to_pdf(docx_bytes_io):
    result = {}

    def run_conversion():
        try:
            pythoncom.CoInitialize()
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as docx_file:
                docx_file.write(docx_bytes_io.getvalue())
                docx_path = docx_file.name

            pdf_path = docx_path.replace(".docx", ".pdf")

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(docx_path)
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close(False)
            word.Quit()

            with open(pdf_path, "rb") as f:
                result['pdf'] = BytesIO(f.read())

        except Exception as e:
            result['error'] = str(e)

        finally:
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

    thread = threading.Thread(target=run_conversion)
    thread.start()
    thread.join()

    if 'error' in result:
        raise Exception(result['error'])

    return result['pdf']

#########################################################################################

####################################################################################################
@app.route('/generate', methods=['POST'])
def generate_proposal():
    data = request.get_json()
    model_type = data.get("modelType", "opex")
    print("Model Type:", model_type)
    print("Incoming Data Keys:", data.keys())

    try:
        if model_type == "capex":
            price = float(data.get("pricePerWp", "0"))
            capacity = float(data.get("capacity", "0").replace(',', ''))
            amount = price * capacity * 1000

            data["price"] = format_inr(price)
            data["amount"] = format_inr(amount)

            doc = Document("proposal-template-CAPEX.docx")

        else:
            selected_years = data.get("template", "").split(",")
            selected_years = [int(y.strip()) for y in selected_years if y.strip().isdigit()]
            if not selected_years:
                return {"error": "No valid template selected"}, 400

            selected_years_str = ",".join([str(y) for y in selected_years])
            template_name = f"proposal-temp-{selected_years_str}"

            if data.get("includeSavings", "").lower() == "yes":
                template_name = f"proposal-temp-{'^J'.join(map(str, selected_years))}-savings"

            if not os.path.exists(template_name + ".docx"):
                return jsonify({"error": f"Template '{template_name}.docx' not found"}), 404

            doc = Document(template_name + ".docx")

            gen_number = float(data.get("generationNumber", "0").replace(',', ''))
            capacity = float(data.get("capacity", "0").replace(',', ''))

            for y in [10, 15, 25]:
                key = f"tariff_{y}"
                val = data.get(key)
                if val:
                    data[f"tariff-{y}"] = str(val)
                    data[f"Tariff-{y}"] = str(val)

            gen_first_year = round(gen_number * capacity)
            data["generationFirstYear"] = format_inr(gen_first_year)

            expected_gen = []
            for i in range(25):
                if i == 0:
                    val = gen_first_year
                elif i == 1 or i == 2:
                    val = expected_gen[-1] * (1 - 0.015)
                else:
                    val = expected_gen[-1] * (1 - 0.0055)
                expected_gen.append(round(val))
                data[f"expectedGen{i+1}"] = format_inr(val)
                data[f"guaranteedGen{i+1}"] = format_inr(val * 0.9)

            full_gen_list = [g / 1e6 for g in expected_gen]
            o_m_base = [0.36 * capacity / 1000]
            for i in range(1, max(selected_years)):
                o_m_base.append(o_m_base[-1] * 1.03)

            for year in selected_years:
                tariff = float(data.get(f"tariff_{year}", "0"))
                revenue_list = [g * tariff for g in full_gen_list[:year]]
                o_m_list = o_m_base[:year]

                for i in range(year):
                    key = f"termPayment-{year}-{i+1}"
                    if i < 3:
                        data[key] = "X"
                    else:
                        npv1 = calculate_npv(revenue_list[i:], 0.15)
                        npv2 = calculate_npv(o_m_list[i:], 0.15)
                        termination = (npv1 - npv2) * 1_000_000 / capacity
                        data[key] = format_inr(termination)

            if data.get("includeSavings", "").lower() == "yes":
                base_grid_price = float(data.get("current_electricity_price-1", "0"))
                for year in selected_years:
                    total_savings = 0
                    solar_tariff = float(data.get(f"tariff_{year}", "0"))
                    for i in range(year):
                        exp_gen = expected_gen[i]
                        grid_price = base_grid_price * ((1.015) ** i)
                        annual_grid_bill = exp_gen * grid_price
                        annual_solar_bill = exp_gen * solar_tariff
                        annual_saving = annual_grid_bill - annual_solar_bill

                        data[f"current_electricity_price-{i+1}"] = "{:.2f}".format(grid_price)
                        data[f"Annual_Grid_Bill-{i+1}"] = format_inr(annual_grid_bill)
                        data[f"Annual_Solar_Bill-{i+1}"] = format_inr(annual_solar_bill)
                        data[f"Annual_Cost_Savings-{i+1}"] = format_inr(annual_saving)
                        total_savings += annual_saving
                    data[f"Sum_Cost_Savings-{year}"] = format_inr(total_savings)

        replace_all_placeholders(doc, data)

        file_format = data.get('format', 'docx').lower()
        filename = f"Proposal_{data.get('clientName', 'Client')}.{file_format}"

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        if file_format == "pdf":
            output = convert_docx_to_pdf(output)
            mimetype = "application/pdf"
        else:
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return send_file(output, as_attachment=True, download_name=filename, mimetype=mimetype)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
